import streamlit as st

def generate_report_html(patient, results):
    """
    Generates a beautifully styled, print-ready HTML medical report.
    Uses clean inline CSS and a print-media query to look professional on paper or PDF.
    """
    p_id = patient.get("patient_id", "N/A")
    name = patient.get("name", "N/A")
    age = patient.get("age", "N/A")
    gender = patient.get("gender", "N/A")
    location = patient.get("location", "N/A")
    history = patient.get("medical_history", "None")
    meds = patient.get("medications", "None")
    fam = patient.get("family_history", "No History")
    neuro = patient.get("neurological_conditions", "None")
    
    score = results.get("final_score", 0.0)
    category = results.get("risk_category", "Low Risk")
    rec = results.get("recommendation", "")
    date = results.get("screening_date", "")
    
    s_score = results.get("symptom_score", 0.0)
    v_score = f"{results.get('voice_score'):.1f}/100" if results.get('voice_score') is not None else "Not Screened"
    w_score = f"{results.get('writing_score'):.1f}/100" if results.get('writing_score') is not None else "Not Screened"
    g_score = f"{results.get('gait_score'):.1f}/100" if results.get('gait_score') is not None else "Not Screened"
    
    # Map badge class
    badge_color = "#065f46" if category == "Low Risk" else "#92400e" if category == "Moderate Risk" else "#991b1b"
    badge_bg = "#d1fae5" if category == "Low Risk" else "#fef3c7" if category == "Moderate Risk" else "#fee2e2"
    
    factors_html = "".join([f"<li>{factor}</li>" for factor in results.get("contributing_factors", [])])
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>ParkiSense - Screening Report</title>
        <style>
            body {{
                font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
                color: #333;
                line-height: 1.5;
                margin: 0;
                padding: 40px;
                background-color: #ffffff;
            }}
            .header {{
                border-bottom: 3px solid #0f766e;
                padding-bottom: 20px;
                margin-bottom: 30px;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }}
            .header h1 {{
                color: #0f766e;
                margin: 0;
                font-size: 26px;
                font-weight: 700;
                letter-spacing: -0.5px;
            }}
            .header .subtitle {{
                color: #64748b;
                font-size: 14px;
                margin-top: 5px;
            }}
            .report-meta {{
                text-align: right;
                font-size: 13px;
                color: #64748b;
            }}
            .section {{
                margin-bottom: 25px;
            }}
            .section-title {{
                font-size: 16px;
                font-weight: 700;
                color: #0f766e;
                border-bottom: 1px solid #e2e8f0;
                padding-bottom: 6px;
                margin-bottom: 12px;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }}
            .grid {{
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 15px;
            }}
            .grid-item {{
                font-size: 14px;
            }}
            .grid-item strong {{
                color: #475569;
            }}
            .score-card {{
                background-color: #f8fafc;
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                padding: 20px;
                margin-bottom: 25px;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }}
            .score-card .score-val {{
                font-size: 36px;
                font-weight: 800;
                color: #0f766e;
                margin: 0;
            }}
            .badge {{
                padding: 8px 16px;
                border-radius: 20px;
                font-weight: 700;
                font-size: 13px;
                color: {badge_color};
                background-color: {badge_bg};
                border: 1px solid {badge_color}33;
                text-transform: uppercase;
            }}
            .table-results {{
                width: 100%;
                border-collapse: collapse;
                margin-top: 10px;
                font-size: 14px;
            }}
            .table-results th, .table-results td {{
                border: 1px solid #e2e8f0;
                padding: 10px;
                text-align: left;
            }}
            .table-results th {{
                background-color: #f8fafc;
                font-weight: 700;
                color: #475569;
            }}
            .disclaimer {{
                background-color: #fafafa;
                border: 1px solid #e5e5e5;
                border-radius: 6px;
                padding: 15px;
                font-size: 11px;
                color: #666;
                margin-top: 40px;
                line-height: 1.4;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
                .no-print {{
                    display: none;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <div>
                <h1>ParkiSense</h1>
                <div class="subtitle">Multimodal Parkinsonian Risk Screening Report</div>
            </div>
            <div class="report-meta">
                <strong>Patient ID:</strong> {p_id}<br>
                <strong>Date:</strong> {date}
            </div>
        </div>
        
        <div class="score-card">
            <div>
                <span style="font-size: 12px; font-weight:700; color:#64748b; text-transform:uppercase;">Overall Parkinsonian Risk Index</span>
                <h2 class="score-val">{score} <span style="font-size: 18px; font-weight:500; color:#64748b;">/ 100</span></h2>
            </div>
            <div>
                <span class="badge">{category}</span>
            </div>
        </div>
        
        <div class="section">
            <div class="section-title">Patient Profile Information</div>
            <div class="grid">
                <div class="grid-item"><strong>Patient Name:</strong> {name}</div>
                <div class="grid-item"><strong>Age / Gender:</strong> {age} years / {gender}</div>
                <div class="grid-item"><strong>Location:</strong> {location}</div>
                <div class="grid-item"><strong>Family History of PD:</strong> {fam}</div>
                <div class="grid-item"><strong>Neurological Profile:</strong> {neuro}</div>
            </div>
            <div style="font-size: 14px; margin-top: 10px;">
                <strong>Clinical Intake History:</strong> {history}<br>
                <strong>Intake Medications:</strong> {meds}
            </div>
        </div>
        
        <div class="section">
            <div class="section-title">Multimodal Screening Matrix</div>
            <table class="table-results">
                <thead>
                    <tr>
                        <th>Assessment Modality</th>
                        <th>Weight</th>
                        <th>Pathological Score</th>
                        <th>Evaluation Method</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>Symptom Checklist</strong></td>
                        <td>40%</td>
                        <td>{s_score:.1f} / 100</td>
                        <td>Clinical Motor & Non-Motor Questionnaire</td>
                    </tr>
                    <tr>
                        <td><strong>Speech & Voice</strong></td>
                        <td>20%</td>
                        <td>{v_score}</td>
                        <td>Sustained Phonation (AH) Frequency Perturbation</td>
                    </tr>
                    <tr>
                        <td><strong>Handwriting & Motor Drawing</strong></td>
                        <td>15%</td>
                        <td>{w_score}</td>
                        <td>Archimedean Spiral Tracing Residual Analysis</td>
                    </tr>
                    <tr>
                        <td><strong>Gait & Lower-Limb Kinematics</strong></td>
                        <td>25%</td>
                        <td>{g_score}</td>
                        <td>CV Spatiotemporal Joint-Angle Estimation</td>
                    </tr>
                </tbody>
            </table>
        </div>
        
        <div class="section">
            <div class="section-title">Key Risk Contributors</div>
            <ul style="font-size: 14px; margin-top: 5px; padding-left: 20px;">
                {factors_html}
            </ul>
        </div>
        
        <div class="section">
            <div class="section-title">Physician Action Plan Recommendations</div>
            <p style="font-size: 14px; margin-top: 5px; color:#0f766e; font-weight:500;">
                {rec}
            </p>
        </div>
        
        <div class="disclaimer">
            <strong>IMPORTANT CLINICAL SAFETY NOTICE:</strong><br>
            This screening report is generated by a prototype AI-assisted risk stratification engine. 
            This is <strong>NOT a clinical medical diagnosis</strong> and does not replace evaluation by a licensed neurologist. 
            All screening findings are preliminary correlation indices based on vocal, motor drawing, and gait kinematic features 
            and must be clinically verified with standard diagnostic procedures (such as MDS-UPDRS examinations or DaTscan imaging).
        </div>
        
        <div class="no-print" style="margin-top: 30px; text-align: center;">
            <button onclick="window.print()" style="background-color: #0f766e; color: white; border: none; padding: 10px 20px; font-size: 14px; font-weight: 500; border-radius: 6px; cursor: pointer;">
                Print Report / Save as PDF
            </button>
        </div>
    </body>
    </html>
    """
    return html

def set_cell_background(cell, fill_hex):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_report_docx(patient, results):
    """
    Generates a professionally designed clinical DOCX report for ParkiSense.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    doc = Document()
    
    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("ParkiSense")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # Teal
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Subtitle
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Multimodal Parkinsonian Risk Screening Report")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate light
    
    # Metadata
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(10)
    meta.paragraph_format.space_after = Pt(20)
    meta.add_run(f"Patient ID: {patient.get('patient_id', 'N/A')}\n")
    meta.add_run(f"Screening Date: {results.get('screening_date', '')}\n")
    
    # Add horizontal rule / divider
    p_hr = doc.add_paragraph()
    p_hr_run = p_hr.add_run("―" * 55)
    p_hr_run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    p_hr.paragraph_format.space_after = Pt(15)
    
    # Score Header
    h_score = doc.add_paragraph()
    h_score_run = h_score.add_run("OVERALL RISK INDEX PROFILE")
    h_score_run.bold = True
    h_score_run.font.size = Pt(14)
    h_score_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    h_score.paragraph_format.space_after = Pt(10)
    
    # Score Summary Box (rendered as a 1x1 table for styling background)
    score_table = doc.add_table(rows=1, cols=1)
    score_table.autofit = False
    score_table.columns[0].width = Inches(6.5)
    cell = score_table.cell(0, 0)
    
    # Set background based on risk
    category = results.get("risk_category", "Low Risk")
    if category == "Low Risk":
        set_cell_background(cell, "D1FAE5") # Light green
        t_color = RGBColor(0x06, 0x5F, 0x46)
    elif category == "Moderate Risk":
        set_cell_background(cell, "FEF3C7") # Light yellow
        t_color = RGBColor(0x92, 0x40, 0x0E)
    else:
        set_cell_background(cell, "FEE2E2") # Light red
        t_color = RGBColor(0x99, 0x1B, 0x1B)
        
    p_cell = cell.paragraphs[0]
    p_cell.paragraph_format.space_before = Pt(8)
    p_cell.paragraph_format.space_after = Pt(8)
    p_cell_run1 = p_cell.add_run("Final Risk Score: ")
    p_cell_run1.font.size = Pt(12)
    p_cell_run2 = p_cell.add_run(f"{results.get('final_score', 0.0)} / 100")
    p_cell_run2.bold = True
    p_cell_run2.font.size = Pt(18)
    p_cell_run2.font.color.rgb = t_color
    p_cell_run3 = p_cell.add_run(f"   |   Risk Classification: {category}")
    p_cell_run3.bold = True
    p_cell_run3.font.size = Pt(12)
    p_cell_run3.font.color.rgb = t_color
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Demographics Section
    h_demo = doc.add_paragraph()
    h_demo_run = h_demo.add_run("PATIENT CLINICAL INTAKE PROFILE")
    h_demo_run.bold = True
    h_demo_run.font.size = Pt(12)
    h_demo_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    h_demo.paragraph_format.space_after = Pt(10)
    
    demo_table = doc.add_table(rows=5, cols=2)
    demo_table.autofit = True
    
    headers = [
        ("Patient Name", patient.get("name", "N/A")),
        ("Age / Gender", f"{patient.get('age', 'N/A')} years / {patient.get('gender', 'N/A')}"),
        ("Location", patient.get("location", "N/A")),
        ("Family History of PD", patient.get("family_history", "No History")),
        ("Neurological Profile", patient.get("neurological_conditions", "None"))
    ]
    for idx, (label, val) in enumerate(headers):
        row = demo_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(str(val))
        set_cell_background(row.cells[0], "F8FAFC")
        
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_hist = doc.add_paragraph()
    p_hist.add_run("Clinical History: ").bold = True
    p_hist.add_run(patient.get("medical_history", "None"))
    
    p_meds = doc.add_paragraph()
    p_meds.add_run("Current Medications: ").bold = True
    p_meds.add_run(patient.get("medications", "None"))
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Multimodal Matrix
    h_matrix = doc.add_paragraph()
    h_matrix_run = h_matrix.add_run("MULTIMODAL SCREENING MATRIX")
    h_matrix_run.bold = True
    h_matrix_run.font.size = Pt(12)
    h_matrix_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    h_matrix.paragraph_format.space_after = Pt(10)
    
    matrix_table = doc.add_table(rows=5, cols=4)
    matrix_table.style = 'Table Grid'
    
    # Headers
    hdr_row = matrix_table.rows[0]
    hdr_cols = ["Assessment Modality", "Weight", "Score Received", "Clinical Methodology"]
    for i, col_name in enumerate(hdr_cols):
        cell = hdr_row.cells[i]
        p = cell.paragraphs[0]
        p_run = p.add_run(col_name)
        p_run.bold = True
        p_run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0F766E")
        
    s_score = results.get("symptom_score", 0.0)
    v_score = f"{results.get('voice_score'):.1f} / 100" if results.get('voice_score') is not None else "Not Screened"
    w_score = f"{results.get('writing_score'):.1f} / 100" if results.get('writing_score') is not None else "Not Screened"
    g_score = f"{results.get('gait_score'):.1f} / 100" if results.get('gait_score') is not None else "Not Screened"
    
    rows_data = [
        ("Symptom Checklist", "40%", f"{s_score:.1f} / 100", "Clinical Motor & Non-Motor Questionnaire"),
        ("Speech & Voice", "20%", v_score, "Sustained Phonation (AH) Frequency Perturbation"),
        ("Handwriting & Tracing", "15%", w_score, "Archimedean Spiral Tracing Residual Analysis"),
        ("Gait & Kinematics", "25%", g_score, "CV Spatiotemporal Joint-Angle Estimation")
    ]
    for idx, data in enumerate(rows_data):
        row = matrix_table.rows[idx+1]
        for col_idx, text in enumerate(data):
            row.cells[col_idx].paragraphs[0].add_run(text)
            if idx % 2 == 1:
                set_cell_background(row.cells[col_idx], "F8FAFC")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Contributing factors
    h_contrib = doc.add_paragraph()
    h_contrib_run = h_contrib.add_run("KEY CONTRIBUTING RISK FACTORS")
    h_contrib_run.bold = True
    h_contrib_run.font.size = Pt(12)
    h_contrib_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    h_contrib.paragraph_format.space_after = Pt(5)
    
    for factor in results.get("contributing_factors", []):
        doc.add_paragraph(f"• {factor}", style='List Bullet')
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Physician Action Plan Recommendation
    h_rec = doc.add_paragraph()
    h_rec_run = h_rec.add_run("CLINICAL RECOMMENDATION ACTION PLAN")
    h_rec_run.bold = True
    h_rec_run.font.size = Pt(12)
    h_rec_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    h_rec.paragraph_format.space_after = Pt(5)
    
    p_rec = doc.add_paragraph()
    p_rec_run = p_rec.add_run(results.get("recommendation", ""))
    p_rec_run.font.size = Pt(11)
    p_rec_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    p_rec_run.bold = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(15)
    
    # Safety Disclaimer Box
    disc_table = doc.add_table(rows=1, cols=1)
    disc_table.autofit = False
    disc_table.columns[0].width = Inches(6.5)
    disc_cell = disc_table.cell(0, 0)
    set_cell_background(disc_cell, "F8FAFC")
    p_disc = disc_cell.paragraphs[0]
    p_disc.paragraph_format.space_before = Pt(6)
    p_disc.paragraph_format.space_after = Pt(6)
    p_disc_run1 = p_disc.add_run("IMPORTANT SAFETY NOTICE: ")
    p_disc_run1.bold = True
    p_disc_run1.font.size = Pt(9)
    p_disc_run1.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p_disc_run2 = p_disc.add_run(
        "This screening report is generated by a prototype AI-assisted risk stratification engine. "
        "This is NOT a clinical medical diagnosis and does not replace evaluation by a licensed neurologist. "
        "All screening findings are preliminary correlation indices based on vocal, motor drawing, and gait kinematic features "
        "and must be clinically verified with standard diagnostic procedures."
    )
    p_disc_run2.font.size = Pt(9)
    p_disc_run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    # Save to binary stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

